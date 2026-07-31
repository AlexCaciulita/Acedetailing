#!/usr/bin/env python3
"""Build the revised Nova Detailing 2026 business plan."""

from __future__ import annotations

from datetime import date, timedelta
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "PLAN-BUSINESS-NOVA-2026-REALIST.docx"
ASSET_DIR = ROOT / "outputs" / "019fb214-2d58-7fc3-b9c0-d5f6971658cb" / "plan-assets"
CHART_PATH = ASSET_DIR / "financial-ramp.png"
LOGO_PATH = ROOT / "public" / "assets" / "logo.png"

FOREST = "2E433B"
FOREST_DARK = "20322C"
SAGE = "DCE6E0"
SAGE_LIGHT = "F2F6F3"
GOLD = "B5A17A"
GOLD_LIGHT = "EFE8D9"
INK = "18231F"
GRAY = "5D6B65"
GRID = "CBD5CF"
RED = "B42318"
RED_LIGHT = "FDECEC"
GREEN = "18794E"
GREEN_LIGHT = "E6F4EC"
WHITE = "FFFFFF"


MONTHS = ["Aug-26", "Sep-26", "Oct-26", "Nov-26", "Dec-26", "Ian-27", "Feb-27", "Mar-27", "Apr-27", "Mai-27", "Iun-27", "Iul-27"]
REVENUE = [26700, 33800, 38600, 45700, 43400, 50500, 55300, 58100, 62900, 65700, 74800, 74800]
OPERATING = [-15719, -10646, -7286, -2213, -3926, 1147, 4507, 6551, 9911, 11955, 18344, 18344]
CASH = [34281, 23635, 16349, 14136, 10210, 11357, 15864, 22415, 32326, 44281, 62625, 80969]
ACTIVE_CLIENTS = [0, 0, 1, 1, 2, 2, 3, 3, 4, 4, 5, 5]


PROSPECTS = [
    ("1", "Speedwell", "Imobiliar", "A", "COO / Office / Procurement", "office@speedwell.be"),
    ("2", "HILS Development", "Imobiliar", "A", "Operațional / Administrativ", "vanzari@hils.ro"),
    ("3", "REDPORT", "Imobiliar", "A", "COO / Office / Procurement", "office@redport.ro"),
    ("4", "Nusco Imobiliara", "Imobiliar", "A", "Administrativ / Procurement", "info@nusco.ro"),
    ("5", "IMPACT Developer & Contractor", "Imobiliar", "A", "Procurement / Fleet", "office@impactsa.ro"),
    ("6", "Genesis Property", "Property management", "A", "COO / Facility / Procurement", "office@genesisproperty.net"),
    ("7", "Graphein", "Inginerie / BIM", "A", "COO / Operations", "office@graphein.ro"),
    ("8", "Leviatan Design", "Inginerie / construcții", "A", "Procurement / Operations", "formular oficial"),
    ("9", "Zitec", "Tehnologie", "A", "Office / Finance / Operations", "contact@zitec.com"),
    ("10", "mindit.io", "Tehnologie", "A", "Office / COO / Finance", "contact@mindit.io"),
    ("11", "One United Properties", "Imobiliar", "B", "Procurement / Fleet", "office@one.ro"),
    ("12", "Forte Partners", "Imobiliar", "B", "COO / Office / Procurement", "formular oficial"),
    ("13", "Cordia România", "Imobiliar", "B", "Administrativ / Operations", "formular / telefon"),
    ("14", "Hagag Development Europe", "Imobiliar", "B", "Office / Procurement", "formular oficial"),
    ("15", "Softbinator Technologies", "Tehnologie", "B", "Office / Operations / CFO", "contact@softbinator.com"),
    ("16", "Connections Consult", "Tehnologie", "B", "Administrativ / Procurement", "office@connectionsconsult.ro"),
    ("17", "Trencadis", "Tehnologie", "B", "COO / Administrativ", "office@trencadis.ro"),
    ("18", "Bittnet Group", "Tehnologie", "B", "Administrativ / CFO", "askformore@bittnet.ro"),
    ("19", "NNDKP", "Servicii juridice", "B", "COO / Office Manager", "office@nndkp.ro"),
    ("20", "Țuca Zbârcea & Asociații", "Servicii juridice", "B", "Office / Administrativ", "formular oficial"),
    ("21", "Mușat & Asociații", "Servicii juridice", "B", "Office / Administrativ", "general@musat.ro"),
    ("22", "Filip & Company", "Servicii juridice", "B", "COO / Office Manager", "office@filipandcompany.com"),
    ("23", "RTPR", "Servicii juridice", "B", "Office / Finance", "office@rtpr.ro"),
    ("24", "PeliPartners", "Servicii juridice", "B", "Office / COO", "office@pelipartners.com"),
    ("25", "Popovici Nițu Stoica & Asociații", "Servicii juridice", "B", "Office / Administrativ", "office@pnsa.ro"),
    ("26", "Schoenherr România", "Servicii juridice", "C", "Office / Procurement", "office.romania@schoenherr.eu"),
    ("27", "Wolf Theiss România", "Servicii juridice", "C", "Office / Procurement", "bucuresti@wolftheiss.com"),
    ("28", "Concelex", "Construcții", "C", "Fleet / Procurement", "office@concelex.ro"),
    ("29", "Bog’Art", "Construcții", "C", "Procurement / Fleet", "office@bogart.ro"),
    ("30", "Reff & Associates", "Servicii juridice", "C", "Office / Procurement", "formular oficial"),
]


DAILY = [
    ("1", "Ofertă", "Blochează pilotul: 3 mașini, o locație, raport 24–48h, remediere separată.", "Fișă ofertă v1"),
    ("2", "Măsurare", "Pornește time-study pe tehnician, boxă, materiale și rework pentru 20 lucrări.", "100% lucrări pontate"),
    ("3", "Documente", "Ofertă, comandă, termeni, GDPR și disclaimer pre-retur.", "Set comercial gata"),
    ("4", "Raport", "Finalizează șablonul foto, severitate, recomandare și accept client.", "Raport Nova v1"),
    ("5", "CRM", "Importă cei 30 de prospecți și verifică primele 10 canale.", "10 conturi validate"),
    ("6", "Research", "Identifică rolurile-țintă și motivul specific pentru primele 10 companii.", "10 mesaje"),
    ("7", "Script", "Testează email, apel și follow-up; elimină jargonul.", "Script v2"),
    ("8", "Outbound", "Trimite primele 10 emailuri corporate; loghează următorul pas.", "10 atingeri"),
    ("9", "Apel", "Sună 5 companii; cere responsabilul și calendarul retururilor.", "5 apeluri"),
    ("10", "Site", "Publică pagina Companii și structura pilotului.", "Pagină validată"),
    ("11", "Outbound", "Trimite al doilea lot de 10 mesaje, adaptat pe segment.", "20 conturi atinse"),
    ("12", "Follow-up", "Revino la lotul 1 cu o întrebare și un CTA de 15 minute.", "≥3 răspunsuri cumulat"),
    ("13", "Discovery", "Colectează flotă, lessor, retur, durere, aprobator și termen.", "1–2 discovery"),
    ("14", "Ofertare", "Trimite oferta pilot în maximum 2 ore, cu două date.", "1 ofertă trimisă"),
    ("15", "Review", "Analizează 50+ atingeri; dacă răspunsul e <5%, schimbă ICP sau mesajul.", "Decizie documentată"),
    ("16", "Outbound", "Următoarele 10 conturi + 5 introduceri calde.", "15 atingeri"),
    ("17", "Simulare", "Simulează inspecția a 3 mașini și măsoară timpul complet.", "Timp real pilot"),
    ("18", "Follow-up", "Revino la ofertele deschise și cere data deciziei.", "2 oferte active"),
    ("19", "Pilot", "Confirmă în scris; colectează ghid lessor și lista VIN.", "Pilot programat"),
    ("20", "Execuție", "Inspectează și fotografiază; nu remedia fără acord.", "3 fișe complete"),
    ("21", "Livrare", "Trimite raportul în 24–48h cu recomandări pe vehicul.", "SLA respectat"),
    ("22", "Remediere", "Prezintă separat preț, timp și risc; cere aprobare.", "Decizie client"),
    ("23", "Dovadă", "Cere feedback și permisiune pentru studiu de caz anonim.", "Feedback utilizabil"),
    ("24", "Cadru", "Propune acord 6 luni cu SLA și prețuri pe intervenție.", "Propunere cadru"),
    ("25", "Pipeline", "Adaugă 20 de companii similare contului care a răspuns.", "50 conturi CRM"),
    ("26", "SOP", "Recepție, fotografii, chei, aprobare, rework și predare.", "SOP v1"),
    ("27", "Marjă", "Calculează lei/oră, lei/zi boxă și contribuția pilotului.", "Unit economics"),
    ("28", "Închidere", "Negociază scopul/volumul, nu standardul; cere semnătură.", "Obiecție finală sau acord"),
    ("29", "Admin", "Testează contract, factură, e-Factura, arhivare și email.", "Flux testat"),
    ("30", "Gate", "Scalează, ajustează sau oprește ipoteza pe KPI-uri.", "1 pilot + 3 oferte sau pivot"),
]


WEEKLY = [
    ("S1", "Ofertă și măsurare", "20 conturi validate; ofertă, raport și time-study gata"),
    ("S2", "Primul outbound", "35–40 atingeri; minimum 2 răspunsuri"),
    ("S3", "Discovery", "2–3 discovery; 1 ofertă pilot"),
    ("S4", "Primul pilot", "1 pilot programat; minimum 80 atingeri cumulat"),
    ("S5", "Dovadă", "Raport la termen; contribuție pilot ≥55%"),
    ("S6", "Al doilea pilot", "2 piloți cumulat; rework <5%"),
    ("S7", "Studiu de caz", "1 studiu anonim; 1 propunere cadru"),
    ("S8", "Primul acord", "1 client cadru semnat"),
    ("S9", "Pipeline 2", "60 conturi; 120 atingeri/lună"),
    ("S10", "Recurență", "Calendar de retururi pe 90 zile"),
    ("S11", "Parteneriate", "5 introduceri calde"),
    ("S12", "Marjă", "Nicio ofertă sub 55% contribuție"),
    ("S13", "Gate 90 zile", "2 clienți activi sau pivot documentat"),
]


MONTHLY = [
    ("L1", "Validare", "1 pilot plătit; ofertă repetabilă; 20 lucrări măsurate"),
    ("L2", "Dovadă", "2 piloți cumulat; studiu de caz; marjă ≥55%"),
    ("L3", "Primul contract", "1 client cadru; retururi vizibile pe 90 zile"),
    ("L4", "Repetabilitate", "Al doilea client activ; raport standard stabil"),
    ("L5", "Capacitate", "Plan retail+B2B pe boxă și tehnician; rework <4%"),
    ("L6", "Break-even", "Rezultat operațional lunar pozitiv; cash rămâne pozitiv"),
    ("L7", "Recurență", "3 clienți activi; forecast pe 8 săptămâni"),
    ("L8", "Referințe", "2 studii de caz; 5 introduceri calde/lună"),
    ("L9", "Scalare controlată", "4 clienți; pipeline ponderat ≥2× ținta"),
    ("L10", "Eficiență", "Lei/oră tehnician +10% față de luna 2"),
    ("L11", "Portofoliu", "Renegociază clienții sub marjă; retrage serviciile slabe"),
    ("L12", "Profitabilitate", "5 clienți; rezultat anual pozitiv; cash buffer refăcut"),
]


SOURCES = [
    ("ONRC", "Corespondență CAEN Rev.2–Rev.3; vechiul 4520 este asociat cu 9531.", "https://www.onrc.ro/documente/anunturi/Corespondenta-CAEN-Rev.2-CAEN-Rev.3.pdf"),
    ("ANAF", "Cota standard de TVA este 21% începând cu 1 august 2025.", "https://static.anaf.ro/static/10/Anaf/AsistentaContribuabili_r/Cotele_de_TVA_09.2025.pdf"),
    ("ANAF", "În 2026, termenul RO e-Factura este de 5 zile lucrătoare de la emitere.", "https://static.anaf.ro/static/3/Ploiesti/20260115111226_comunicat%20ajfp%20arges%20-%20modificari%20ro%20e-factura%20site.pdf"),
    ("Ayvens România", "Ghidul de retur arată standarde și rolul evaluării de retur.", "https://www.ayvens.com/-/media/ayvens/public/ro/ghid-returnare-vehicule-flota-ex-lp.pdf?rev=-1"),
    ("Business Lease", "Procesul de returnare și evaluare rămâne în controlul lessorului/evaluatorului.", "https://www.businesslease.ro/pentru-clienti/returnarea-masinii"),
    ("EUR-Lex", "GDPR pentru datele colectate prin site, CRM și rapoarte.", "https://eur-lex.europa.eu/legal-content/RO/TXT/?uri=CELEX%3A32016R0679"),
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=GRID, size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=65, start=80, bottom=65, end=80):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_col_widths(table, widths):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Pagina ")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(GRAY)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def add_hyperlink(paragraph, text, url, color=FOREST, underline=True):
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    new_run = OxmlElement("w:r")
    run_props = OxmlElement("w:rPr")
    color_node = OxmlElement("w:color")
    color_node.set(qn("w:val"), color)
    run_props.append(color_node)
    if underline:
        underline_node = OxmlElement("w:u")
        underline_node.set(qn("w:val"), "single")
        run_props.append(underline_node)
    new_run.append(run_props)
    text_node = OxmlElement("w:t")
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def setup_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name in ("Title", "Heading 1", "Heading 2", "Heading 3"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.color.rgb = RGBColor.from_string(FOREST_DARK)
        style.font.bold = True
        style.paragraph_format.keep_with_next = True

    doc.styles["Title"].font.size = Pt(31)
    doc.styles["Heading 1"].font.size = Pt(20)
    doc.styles["Heading 1"].paragraph_format.space_before = Pt(8)
    doc.styles["Heading 1"].paragraph_format.space_after = Pt(8)
    doc.styles["Heading 1"].paragraph_format.page_break_before = True
    doc.styles["Heading 2"].font.size = Pt(14)
    doc.styles["Heading 2"].font.color.rgb = RGBColor.from_string(FOREST)
    doc.styles["Heading 2"].paragraph_format.space_before = Pt(9)
    doc.styles["Heading 2"].paragraph_format.space_after = Pt(5)
    doc.styles["Heading 3"].font.size = Pt(11)
    doc.styles["Heading 3"].font.color.rgb = RGBColor.from_string(GOLD)

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(10.2)
        style.paragraph_format.left_indent = Cm(0.6)
        style.paragraph_format.first_line_indent = Cm(-0.25)
        style.paragraph_format.space_after = Pt(3)


def configure_sections(doc):
    for index, section in enumerate(doc.sections):
        section.top_margin = Inches(0.72)
        section.bottom_margin = Inches(0.65)
        section.left_margin = Inches(0.78)
        section.right_margin = Inches(0.78)
        section.header_distance = Inches(0.25)
        section.footer_distance = Inches(0.28)
        section.different_first_page_header_footer = index == 0


def add_header_footer(section):
    header = section.header
    table = header.add_table(rows=1, cols=2, width=Inches(6.9))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(4.8)
    table.columns[1].width = Inches(2.1)
    left = table.cell(0, 0)
    right = table.cell(0, 1)
    left.text = "NOVA DETAILING"
    right.text = "PLAN DE BUSINESS · 2026"
    for cell in (left, right):
        set_cell_shading(cell, FOREST_DARK)
        set_cell_margins(cell, 45, 90, 45, 90)
        for run in cell.paragraphs[0].runs:
            run.font.size = Pt(8)
            run.font.bold = True
            run.font.color.rgb = RGBColor.from_string(WHITE)
    right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    footer = section.footer
    p = footer.paragraphs[0]
    p.add_run("Confidențial · Nova Detailing · Versiune 31.07.2026")
    p.runs[0].font.size = Pt(8)
    p.runs[0].font.color.rgb = RGBColor.from_string(GRAY)
    add_page_field(footer.add_paragraph())


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("NOVA DETAILING")
    r.font.name = "Calibri"
    r.font.size = Pt(28)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(GOLD)
    p.space_after = Pt(8)

    band = doc.add_table(rows=1, cols=1)
    band.alignment = WD_TABLE_ALIGNMENT.CENTER
    band.autofit = False
    band.cell(0, 0).width = Inches(6.7)
    set_cell_shading(band.cell(0, 0), FOREST_DARK)
    set_cell_margins(band.cell(0, 0), 130, 180, 130, 180)
    title_p = band.cell(0, 0).paragraphs[0]
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("PLAN DE BUSINESS\nREALIST ȘI EXECUTABIL")
    run.font.name = "Calibri"
    run.font.size = Pt(29)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(WHITE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Retail premium + pilot B2B pentru retururi de leasing")
    r.font.size = Pt(17)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(FOREST)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("București–Ilfov · August 2026 – Iulie 2027")
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string(GRAY)

    doc.add_paragraph("")
    info = doc.add_table(rows=4, cols=2)
    info.alignment = WD_TABLE_ALIGNMENT.CENTER
    info.autofit = False
    rows = [
        ("Document", "Plan refăcut după analiza operațională și B2B"),
        ("Obiectiv", "Primul client în 30–45 zile; model lunar profitabil în jurul lunii 6"),
        ("Bază financiară", "Scenariu net de TVA, cu costuri fixe și variabile explicite"),
        ("Statut", "Ipoteze de lucru — se înlocuiesc cu date reale după 20 de lucrări"),
    ]
    for i, values in enumerate(rows):
        for j, value in enumerate(values):
            cell = info.cell(i, j)
            cell.text = value
            set_cell_border(cell)
            set_cell_margins(cell, 95, 110, 95, 110)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if j == 0:
                set_cell_shading(cell, GOLD_LIGHT)
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string(FOREST_DARK)
    set_col_widths(info, [Inches(1.55), Inches(5.15)])

    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("CONFIDENȚIAL · DOCUMENT DE LUCRU")
    r.font.size = Pt(9)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(GOLD)


def add_table(doc, headers, rows, widths, font_size=8.5, header_fill=FOREST, repeat_header=True):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    header = table.rows[0]
    if repeat_header:
        set_repeat_table_header(header)
    for idx, text in enumerate(headers):
        cell = header.cells[idx]
        cell.text = str(text)
        set_cell_shading(cell, header_fill)
        set_cell_border(cell)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(font_size)
                run.font.color.rgb = RGBColor.from_string(WHITE)
    header.height = Cm(0.65)
    header.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

    for ridx, row_data in enumerate(rows):
        row = table.add_row()
        prevent_row_split(row)
        for idx, value in enumerate(row_data):
            cell = row.cells[idx]
            cell.text = str(value)
            set_cell_border(cell)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if ridx % 2 == 1:
                set_cell_shading(cell, SAGE_LIGHT)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(font_size)
                    run.font.color.rgb = RGBColor.from_string(INK)
    set_col_widths(table, widths)
    return table


def add_callout(doc, title, body, tone="gold"):
    fill = GOLD_LIGHT if tone == "gold" else (GREEN_LIGHT if tone == "green" else RED_LIGHT)
    accent = GOLD if tone == "gold" else (GREEN if tone == "green" else RED)
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.cell(0, 0).width = Inches(0.12)
    table.cell(0, 1).width = Inches(6.55)
    set_cell_shading(table.cell(0, 0), accent)
    set_cell_shading(table.cell(0, 1), fill)
    set_cell_margins(table.cell(0, 1), 105, 130, 105, 130)
    p = table.cell(0, 1).paragraphs[0]
    run = p.add_run(title + "\n")
    run.font.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor.from_string(FOREST_DARK)
    run = p.add_run(body)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_bullets(doc, items, numbered=False):
    style = "List Number" if numbered else "List Bullet"
    for item in items:
        p = doc.add_paragraph(style=style)
        p.add_run(item)


def add_kpi_cards(doc, cards):
    table = doc.add_table(rows=1, cols=len(cards))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for idx, (number, label) in enumerate(cards):
        cell = table.cell(0, idx)
        set_cell_shading(cell, FOREST_DARK if idx % 2 == 0 else FOREST)
        set_cell_margins(cell, 110, 80, 110, 80)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(number + "\n")
        r.font.size = Pt(19)
        r.font.bold = True
        r.font.color.rgb = RGBColor.from_string(GOLD)
        r = p.add_run(label)
        r.font.size = Pt(8.4)
        r.font.bold = True
        r.font.color.rgb = RGBColor.from_string(WHITE)
    set_col_widths(table, [Inches(6.7 / len(cards))] * len(cards))


def business_date(day_index):
    current = date(2026, 8, 3)
    count = 1
    while count < day_index:
        current += timedelta(days=1)
        if current.weekday() < 5:
            count += 1
    return current.strftime("%d.%m")


def create_chart():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    width, height = 1800, 760
    image = Image.new("RGB", (width, height), "#F7F8F6")
    draw = ImageDraw.Draw(image)
    regular_path = "/System/Library/Fonts/Supplemental/Arial.ttf"
    bold_path = "/System/Library/Fonts/Supplemental/Arial Bold.ttf"
    regular = ImageFont.truetype(regular_path, 24)
    small = ImageFont.truetype(regular_path, 20)
    title_font = ImageFont.truetype(bold_path, 38)
    legend_font = ImageFont.truetype(bold_path, 21)

    left, top, right, bottom = 125, 150, 1740, 650
    draw.text((left, 42), "Rampa financiară — scenariu de bază", fill=f"#{FOREST_DARK}", font=title_font)

    min_value, max_value = -20000, 80000
    ticks = [-20000, 0, 20000, 40000, 60000, 80000]
    for tick in ticks:
        y = bottom - int((tick - min_value) / (max_value - min_value) * (bottom - top))
        draw.line((left, y, right, y), fill="#DDE3DF", width=2)
        draw.text((18, y - 12), f"{tick // 1000}k", fill="#5D6B65", font=small)

    def point(index, value):
        x = left + int(index / (len(MONTHS) - 1) * (right - left))
        y = bottom - int((value - min_value) / (max_value - min_value) * (bottom - top))
        return x, y

    zero_y = point(0, 0)[1]
    draw.line((left, zero_y, right, zero_y), fill="#8A9690", width=3)
    revenue_points = [point(i, value) for i, value in enumerate(REVENUE)]
    operating_points = [point(i, value) for i, value in enumerate(OPERATING)]
    draw.line(revenue_points, fill=f"#{FOREST}", width=7, joint="curve")
    draw.line(operating_points, fill=f"#{GOLD}", width=7, joint="curve")
    for x, y in revenue_points:
        draw.ellipse((x - 7, y - 7, x + 7, y + 7), fill=f"#{FOREST}")
    for x, y in operating_points:
        draw.ellipse((x - 7, y - 7, x + 7, y + 7), fill=f"#{GOLD}")
    for i, month in enumerate(MONTHS):
        x, _ = point(i, 0)
        bbox = draw.textbbox((0, 0), month, font=small)
        draw.text((x - (bbox[2] - bbox[0]) / 2, bottom + 25), month, fill="#46544E", font=small)

    draw.line((left, 105, left + 55, 105), fill=f"#{FOREST}", width=7)
    draw.text((left + 70, 91), "Venit net", fill=f"#{FOREST_DARK}", font=legend_font)
    draw.line((left + 300, 105, left + 355, 105), fill=f"#{GOLD}", width=7)
    draw.text((left + 370, 91), "Rezultat operațional", fill=f"#{FOREST_DARK}", font=legend_font)
    image.save(CHART_PATH, quality=94)


def build():
    create_chart()
    doc = Document()
    setup_styles(doc)
    configure_sections(doc)
    add_cover(doc)

    doc.add_heading("1. Decizia executivă", level=1)
    add_callout(
        doc,
        "Teza de business",
        "Nova nu trebuie să se transforme într-o spălătorie de flotă. Modelul recomandat păstrează retailul premium ca motor de marjă și adaugă un serviciu B2B îngust: pilot pre-retur pentru 3 autoturisme, cu inspecție foto, raport în 24–48 de ore și remedieri aprobate separat.",
        "green",
    )
    add_kpi_cards(doc, [
        ("30–45 zile", "țintă pentru primul client"),
        ("3 mașini", "pilotul de intrare"),
        ("≥55%", "contribuție minimă pilot"),
        ("Ian. 2027", "prima lună pozitivă în bază"),
    ])
    doc.add_paragraph("")
    doc.add_heading("Ce trebuie să fie adevărat", level=2)
    add_bullets(doc, [
        "Retailul menține 8–12 lucrări/lună la un ticket mediu net de aproximativ 2.800 lei.",
        "Piloții cresc de la 1 la 3–4/lună, iar până în luna 5 există minimum 2 clienți recurenți.",
        "Costul fix lunar rămâne în jur de 35.100 lei și include remunerația fondatorului și doi tehnicieni.",
        "Capacitatea este condusă pe zile-boxă, nu pe numărul brut de programări; pragul utilizabil este 85%.",
        "Fiecare preț este recalculat după 20 de lucrări măsurate, pe lei/oră tehnician și lei/zi boxă.",
    ])
    doc.add_heading("Rezultatul scenariului de bază", level=2)
    add_table(
        doc,
        ["Indicator", "Valoare", "Interpretare"],
        [
            ("Venit net 12 luni", "630.300 lei", "Fără TVA; fără venituri din școală"),
            ("Rezultat operațional", "30.969 lei", "Înainte de impozit, datorii și investiții majore"),
            ("Cash inițial", "50.000 lei", "Ipoteză — de înlocuit cu soldul real"),
            ("Cash minim", "10.210 lei", "Risc în luna 5"),
            ("Capital suplimentar pentru buffer", "34.790 lei", "Necesar pentru a păstra bufferul de 45.000 lei"),
            ("Clienți B2B la final", "5", "Obținuți gradual; nu sunt contractați astăzi"),
        ],
        [Inches(2.0), Inches(1.35), Inches(3.35)],
        9,
    )
    add_callout(
        doc,
        "Concluzie",
        "Modelul poate deveni profitabil lunar în jurul lunii 6, dar este subcapitalizat față de bufferul recomandat. Până când bufferul este finanțat, sunt blocate angajările suplimentare, capex-ul mare și campaniile plătite agresive.",
        "red",
    )
    doc.add_heading("2. Punctul de plecare și corecțiile obligatorii", level=1)
    doc.add_heading("Ce există deja", level=2)
    add_bullets(doc, [
        "Un atelier poziționat premium, pachete retail, procese de rezervare și o școală de detailing.",
        "Un site nou ca direcție vizuală, dar cu diferențe între codul local și deploymentul public.",
        "Un raport foto Nova și o idee B2B valoroasă, dar încă nevalidată prin piloți plătiți.",
        "Active comerciale utile: experiență declarată, portofoliu, procese de atelier și servicii cu ticket ridicat.",
    ])
    doc.add_heading("Ce se oprește imediat", level=2)
    add_table(
        doc,
        ["Element", "Problemă", "Decizie"],
        [
            ("VIP anual 6.000–9.000 lei", "Economia pachetului nu susține 2 Signature + 12 mentenanțe.", "Se retrage; se oferă doar mentenanță personalizată după evaluare."),
            ("Promisiunea avansului online", "Plata nu trebuie promisă până când PayU și emailul sunt configurate și testate.", "CTA neutru: programare confirmată după verificare."),
            ("Durate fixe contradictorii", "Premium/Signature au durate diferite între pagini.", "Durata se confirmă după inspecție până la time-study."),
            ("„Raport oficial de retur”", "Evaluarea finală aparține lessorului/evaluatorului.", "Raportul Nova este diagnostic pre-retur, cu disclaimer clar."),
            ("Creștere prin discount", "Consumă boxă și tehnician fără să protejeze contribuția.", "Discount doar pentru scop redus sau volum ferm."),
        ],
        [Inches(1.55), Inches(2.55), Inches(2.6)],
        8.6,
    )
    doc.add_heading("Ordinea corectă a priorităților", level=2)
    add_bullets(doc, [
        "1) măsurare și ofertă; 2) primul pilot; 3) contract cadru; 4) repetabilitate; 5) scalare.",
        "Școala rămâne o linie opțională. Nu intră în scenariul financiar de bază până când cererea și marja sunt validate separat.",
        "Marketingul plătit nu compensează o ofertă B2B neclară sau un flux tehnic netestat.",
    ])
    doc.add_heading("3. Modelul de business revizuit", level=1)
    doc.add_heading("Arhitectură simplă, cu trei motoare", level=2)
    add_table(
        doc,
        ["Motor", "Rol", "Client", "Venit", "Regulă"],
        [
            ("Retail premium", "Cash și marjă", "Proprietar auto", "Ticket net țintă ~2.800 lei", "Protejează sloturile și prețul"),
            ("Pilot B2B", "Achiziție și dovadă", "Companie cu 5–30 autoturisme", "1.350–1.800 lei + TVA / 3 mașini", "Raport 24–48h; remedieri separat"),
            ("B2B recurent", "Predictibilitate", "Client pilot convertit", "~4.800 lei net/client/lună în bază", "SLA, prețuri și calendar de retur"),
            ("Școală", "Opțiune, nu bază", "Cursant", "Se modelează separat", "Nu finanțează artificial operaționalul"),
        ],
        [Inches(1.25), Inches(1.25), Inches(1.55), Inches(1.45), Inches(1.9)],
        8.4,
    )
    doc.add_heading("Clientul ideal inițial", level=2)
    add_bullets(doc, [
        "Companie owner-led sau filială locală, București–Ilfov, cu 5–30 autoturisme de pasageri.",
        "Are 2–5 retururi de leasing în următoarele 90 de zile și nu are un proces intern riguros de pre-retur.",
        "Decident accesibil: Office Manager, Operations, Procurement, Administrativ, Fleet sau CFO.",
        "Poate aproba un pilot de 3 mașini fără licitație regională sau onboarding de luni de zile.",
    ])
    doc.add_heading("Segmente în ordinea priorității", level=2)
    add_table(
        doc,
        ["Prioritate", "Segment", "De ce", "Risc"],
        [
            ("A", "Dezvoltatori imobiliari și inginerie", "Deplasări locale, proiecte, management accesibil.", "Flota trebuie confirmată."),
            ("A/B", "Tehnologie antreprenorială", "Decizie locală și cultură de pilot.", "Unele echipe sunt remote și fără flotă."),
            ("B", "Firme de avocatură", "Vehicule premium și grijă pentru discreție.", "Lot mic; sensibilitate la contact."),
            ("C", "Contractori mari", "Flotă probabilă și valoare mare.", "Procurement lung; utilitarele nu sunt ICP inițial."),
        ],
        [Inches(0.75), Inches(2.0), Inches(2.35), Inches(1.6)],
        8.6,
    )
    add_callout(doc, "Regulă de calificare", "Nu se trimite ofertă completă înainte de a confirma: număr de autoturisme, lessor, date de retur, aprobator, problemă și termen.", "gold")
    doc.add_heading("4. Oferta B2B: „Pilot Retur Fără Surprize”", level=1)
    doc.add_heading("Ce primește clientul", level=2)
    add_bullets(doc, [
        "Inspecție vizuală documentată pentru 3 autoturisme, într-o singură locație din București–Ilfov.",
        "Set foto standard și raport pe fiecare vehicul: constatare, severitate, recomandare și prioritate.",
        "Raport agregat de decizie livrat în 24–48 de ore după inspecție.",
        "Deviz separat pentru remedieri; nicio lucrare nu începe fără aprobare scrisă.",
        "Opțional, predare la atelier și re-inspecție internă după remediere.",
    ])
    doc.add_heading("Ce nu promite", level=2)
    add_bullets(doc, [
        "Nu înlocuiește evaluarea oficială a companiei de leasing sau a evaluatorului desemnat.",
        "Nu garantează eliminarea tuturor costurilor de retur.",
        "Nu oferă un verdict juridic și nu modifică ghidul de uzură al lessorului.",
        "Nu include mecanică, tinichigerie structurală sau daune în afara competenței Nova.",
    ])
    doc.add_heading("Preț și limite", level=2)
    add_table(
        doc,
        ["Element", "Pilot", "După validare"],
        [
            ("Inspecție + raport", "1.350–1.800 lei + TVA / 3 mașini", "Preț fix pe lot sau per mașină"),
            ("Remedieri", "Separate, aprobate pe vehicul", "Grilă pe tip de intervenție"),
            ("SLA", "24–48h pentru raport", "SLA contractual doar după 2 piloți reușiți"),
            ("Deplasare", "O locație inclusă", "Tarif zonal / kilometraj"),
            ("Plată", "Conform ofertei și bonității", "Termen negociat; e-Factura în flux"),
        ],
        [Inches(1.7), Inches(2.5), Inches(2.5)],
        8.7,
    )
    doc.add_heading("Contractul cadru după pilot", level=2)
    add_bullets(doc, [
        "Durată inițială de 6 luni, fără exclusivitate și fără promisiuni de volum neconfirmate.",
        "Anexă de servicii și prețuri, SLA, procedură de aprobare, custodie chei, răspundere și protecția datelor.",
        "Calendar de retururi pe 90 zile, actualizat lunar de client.",
        "Revizuire de preț dacă mixul sau timpul real diferă cu peste 15% față de pilot.",
    ])
    doc.add_heading("5. Cum se semnează primul client", level=1)
    add_kpi_cards(doc, [
        ("50–60", "companii în listă"),
        ("100", "contacte/canale"),
        ("80–120", "atingeri"),
        ("5 / 3 / 1", "discovery / oferte / client"),
    ])
    doc.add_paragraph("")
    doc.add_heading("Funnel și ritm", level=2)
    add_table(
        doc,
        ["Etapă", "Țintă", "Definiție", "Termen"],
        [
            ("Conturi", "40/lună", "Fit geografic și rol-țintă", "10/săptămână"),
            ("Contacte", "80/lună", "Minimum 2 canale/roluri per cont", "Înainte de outreach"),
            ("Atingeri", "120/lună", "Email + apel + follow-up", "30/săptămână"),
            ("Răspunsuri", "12/lună", "Răspuns uman relevant", "Optimizare dacă <5%"),
            ("Discovery", "5/lună", "Flotă, retur, lessor, aprobator", "15–20 minute"),
            ("Piloți ofertați", "3/lună", "Ofertă scrisă cu dată", "În 2 ore de la call"),
            ("Piloți executați", "1–2/lună", "Pilot plătit + raport", "SLA 24–48h"),
            ("Contracte", "1/lună după validare", "Acord semnat + fereastră", "Follow-up cu dată"),
        ],
        [Inches(1.35), Inches(1.15), Inches(2.9), Inches(1.3)],
        8.4,
    )
    doc.add_heading("Secvența de contact — 12 zile", level=2)
    add_table(
        doc,
        ["Zi", "Canal", "Mesaj"],
        [
            ("0", "Email", "Întrebare despre 2–5 retururi în următoarele 90 zile."),
            ("1", "Telefon", "Identifică responsabilul și verifică fereastra de retur."),
            ("3", "Email", "Trimite structura pilotului, nu o prezentare lungă."),
            ("6", "Telefon / WhatsApp business", "Clarifică blocajul și cere 15 minute."),
            ("9", "Email", "Exemplu anonim de raport + două date de pilot."),
            ("12", "Break-up", "Închide politicos și cere luna potrivită de revenire."),
        ],
        [Inches(0.55), Inches(1.65), Inches(4.5)],
        8.7,
    )
    add_callout(doc, "Cea mai importantă întrebare", "„Aveți 2–5 autoturisme care ies din leasing în următoarele 90 de zile și cine gestionează procesul?”", "green")
    doc.add_heading("6. Lista de 30 de prospecți", level=1)
    p = doc.add_paragraph()
    p.add_run("Limitare importantă: ").bold = True
    p.add_run("prezența unei flote și calendarul retururilor nu sunt confirmate public. Lista ordonează probabilitatea de fit; fiecare companie se califică înainte de ofertă. Canalele sunt corporate/publice.")
    for block_start in range(0, 30, 10):
        if block_start:
            block_heading = doc.add_heading(f"6.{block_start // 10 + 1}. Prospecți {block_start + 1}–{block_start + 10}", level=2)
            block_heading.paragraph_format.page_break_before = True
        rows = PROSPECTS[block_start:block_start + 10]
        add_table(
            doc,
            ["#", "Companie", "Segment", "P", "Rol-țintă", "Canal public"],
            rows,
            [Inches(0.35), Inches(1.65), Inches(1.2), Inches(0.35), Inches(1.65), Inches(1.55)],
            7.7,
        )
        if block_start == 0:
            doc.add_heading("Cum se folosesc primele 10", level=2)
            add_bullets(doc, [
                "Se abordează în două loturi de câte 5, cu mesaj personalizat pe proiecte și operațiuni.",
                "În primul apel se cere responsabilul și data returului; nu se încearcă vânzarea completă.",
                "Dacă o companie nu are 2–5 retururi apropiate, se setează o dată de revenire și se iese din pipeline activ.",
            ])
    doc.add_heading("7. Mesajele comerciale", level=1)
    doc.add_heading("Email inițial", level=2)
    add_callout(
        doc,
        "Subiect: Aveți 3 mașini care ies din leasing în următoarele 90 zile?",
        "Bună ziua,\n\nNova Detailing ajută echipele din București–Ilfov să verifice vizual autoturismele înainte de retur: 3 mașini, fotografii standard și raport în 24–48h. Orice remediere se propune separat și începe doar cu acordul companiei. Raportul este diagnostic pre-retur și nu înlocuiește evaluarea lessorului.\n\nCine gestionează la [Companie] mașinile care ies din leasing și calendarul retururilor? Dacă există 2–5 retururi în următoarele 90 zile, putem discuta 15 minute.",
        "gold",
    )
    doc.add_heading("Deschidere de apel", level=2)
    add_callout(doc, "Script de 25 secunde", "„Bună ziua, sunt [nume] de la Nova Detailing. Nu sun pentru spălare de flotă. Avem un pilot pentru 3 autoturisme care urmează să fie returnate din leasing: inspecție foto și raport pre-retur. Cine gestionează la dumneavoastră retururile și datele acestora?”", "green")
    objections_heading = doc.add_heading("Obiecții", level=2)
    objections_heading.paragraph_format.page_break_before = True
    add_table(
        doc,
        ["Obiecție", "Răspuns recomandat"],
        [
            ("„Avem deja furnizor.”", "Perfect. Pilotul poate rămâne un al doilea control înainte de retur; nu cere schimbarea furnizorului."),
            ("„Trimiteți o ofertă.”", "O trimit astăzi. Ca să fie relevantă: câte mașini, ce lessor și ce dată de retur?"),
            ("„Nu vrem cost suplimentar.”", "Scopul pilotului este o decizie informată. Remedierea nu este automată și nu se promite economie fără dovadă."),
            ("„Este evaluare oficială?”", "Nu. Este diagnostic pre-retur; evaluatorul/lessorul păstrează decizia finală."),
            ("„Prețul este mare.”", "Putem reduce scopul sau volumul, nu standardul. Comparăm costul pilotului cu timpul și riscul deciziei în orb."),
            ("„Nu avem retururi acum.”", "În ce lună apar următoarele 2–3 retururi? Revin cu 30 zile înainte."),
        ],
        [Inches(1.75), Inches(4.95)],
        8.8,
    )
    doc.add_heading("Reguli de închidere", level=2)
    add_bullets(doc, [
        "Fiecare discuție se termină cu o dată, un proprietar și un pas: nu cu „mai vorbim”.",
        "Oferta pleacă în maximum 2 ore după discovery și are două ferestre de programare.",
        "Follow-up-ul cere decizia sau blocajul, nu „ați văzut emailul?”.",
        "Contractul se cere după pilot, când există dovadă de livrare și marjă reală.",
    ])
    doc.add_heading("8. Modelul operațional", level=1)
    doc.add_heading("Flux standard pentru un pilot", level=2)
    add_table(
        doc,
        ["Pas", "Owner", "Output", "Control"],
        [
            ("1. Calificare", "Fondator", "Flotă, retur, lessor, aprobator", "Fără ofertă dacă lipsesc datele"),
            ("2. Confirmare", "Fondator", "Comandă + VIN + locație + ghid", "Acord scris"),
            ("3. Inspecție", "Tehnician", "Cadre foto și constatări", "Checklist complet"),
            ("4. Triere", "Lead tehnic", "Severitate și recomandare", "Fără promisiuni de taxare"),
            ("5. Raport", "Fondator / QA", "Raport per vehicul + agregat", "Livrare 24–48h"),
            ("6. Deviz", "Fondator", "Preț/timp/risc pe remediere", "Aprobare per mașină"),
            ("7. Execuție", "Atelier", "Lucrare + time-study", "Rework <5%"),
            ("8. Predare", "Fondator", "Raport final + factură", "Arhivare și e-Factura"),
            ("9. Close", "Fondator", "Feedback + propunere cadru", "Dată de decizie"),
        ],
        [Inches(1.25), Inches(1.15), Inches(2.7), Inches(1.6)],
        8.4,
    )
    doc.add_heading("Roluri minime", level=2)
    add_table(
        doc,
        ["Rol", "Responsabilități", "KPI"],
        [
            ("Fondator / GM", "Vânzări, ofertare, QA raport, cash, relația cu clientul", "30 atingeri/săptămână; marjă; cash"),
            ("Lead tehnic", "Standard foto, triere, QA, time-study", "SLA; rework; lei/oră"),
            ("Tehnicieni", "Execuție și pontaj disciplinat", "Timp standard; consumabile; rework"),
            ("Contabil extern", "Facturare, e-Factura, fiscalitate", "0 întârzieri; reconciliere"),
            ("Avocat / consultant", "Contracte, răspundere, GDPR, mediu", "Documente aprobate înainte de scalare"),
        ],
        [Inches(1.35), Inches(3.95), Inches(1.4)],
        8.5,
    )
    doc.add_heading("Rutina de management", level=2)
    add_bullets(doc, [
        "Zilnic, 08:30: capacitate, lucrări, blocaje; 17:30: ore, consumabile și următorul pas comercial.",
        "Luni: pipeline și listă; marți: outbound; miercuri: discovery/piloți; joi: follow-up/oferte; vineri: cash, KPI și QA.",
        "Lunar: P&L, cashflow, capacitate, marjă pe serviciu, clienți sub prag și următoarele 90 zile.",
    ])
    doc.add_heading("9. Economia unitară și planul financiar", level=1)
    add_callout(doc, "Toate valorile sunt ipoteze", "Înlocuiește prețurile, timpii și costurile după 20 de lucrări măsurate. Modelul exclude TVA, impozitul pe profit/micro, ratele, investițiile majore și venitul din școală.", "gold")
    doc.add_heading("Ipoteze principale", level=2)
    add_table(
        doc,
        ["Categorie", "Ipoteză", "Valoare"],
        [
            ("Retail", "Ticket net / zile-boxă / cost variabil", "2.800 lei / 1,8 / 27%"),
            ("Pilot", "Preț / mașini / cost variabil", "1.500 lei / 3 / 25%"),
            ("Remediere", "Ticket / conversie / cost variabil", "2.800 lei / 35% / 32%"),
            ("Recurent", "Venit/client/lună / cost variabil", "4.800 lei / 30%"),
            ("Capacitate", "2 boxe × 22 zile × prag utilizabil", "37,4 zile-boxă la 85%"),
            ("Cost fix", "Fondator, 2 tehnicieni, chirie, utilități, vânzări, admin", "35.100 lei/lună"),
        ],
        [Inches(1.25), Inches(3.5), Inches(1.95)],
        8.8,
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    chart = p.add_run().add_picture(str(CHART_PATH), width=Inches(6.65))
    chart._inline.docPr.set("descr", "Grafic cu venitul net și rezultatul operațional lunar în scenariul financiar de bază.")
    chart._inline.docPr.set("title", "Rampa financiară Nova Detailing")
    caption = doc.add_paragraph("Figura 1. Venit net și rezultat operațional lunar — scenariu de bază.", style="Caption")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    projection_heading = doc.add_heading("9.1. Proiecția lunară", level=2)
    projection_heading.paragraph_format.page_break_before = True
    financial_rows = []
    for idx in range(12):
        financial_rows.append((
            MONTHS[idx],
            f"{REVENUE[idx]:,}".replace(",", "."),
            f"{OPERATING[idx]:,}".replace(",", "."),
            f"{CASH[idx]:,}".replace(",", "."),
            ACTIVE_CLIENTS[idx],
        ))
    add_table(
        doc,
        ["Lună", "Venit net", "Rezultat op.", "Cash final", "Clienți B2B"],
        financial_rows,
        [Inches(0.85), Inches(1.5), Inches(1.5), Inches(1.5), Inches(1.35)],
        8.8,
    )
    doc.add_heading("Interpretare financiară", level=2)
    add_bullets(doc, [
        "Pierderile cumulate ating aproximativ 39.790 lei înainte de revenire; cash-ul minim este aproximativ 10.210 lei.",
        "Break-even-ul lunar apare în ianuarie 2027, când scenariul are 2 clienți recurenți și 3 piloți/lună.",
        "Profitul anual de aproximativ 30.969 lei este modest: modelul nu suportă abatere mare de la vânzări sau costuri.",
        "Pentru a nu coborî sub bufferul de 45.000 lei ar fi necesar cash inițial de aproximativ 84.790 lei sau o reducere echivalentă a pierderilor de ramp-up.",
        "Nu se distribuie dividende și nu se aprobă capex mare până când cash-ul rămâne peste buffer trei luni consecutiv.",
    ])
    doc.add_heading("Praguri de preț", level=2)
    add_table(
        doc,
        ["Măsură", "Prag", "Acțiune dacă nu este atins"],
        [
            ("Contribuție pilot", "≥55%", "Crește prețul sau reduce scopul"),
            ("Contribuție totală", "≥68%", "Reprețuiește mixul și consumabilele"),
            ("Rework", "<5%", "Stop scalare; analiză cauză"),
            ("Utilizare maximă boxe", "≤85%", "Replanifică; nu suprarezerva"),
            ("Cash", "≥45.000 lei buffer", "Blochează capex și costuri discreționare"),
        ],
        [Inches(2.0), Inches(1.4), Inches(3.3)],
        8.8,
    )
    doc.add_heading("10. Plan pe 30 de zile lucrătoare", level=1)
    for block_start in range(0, 30, 10):
        if block_start:
            block_heading = doc.add_heading(f"10.{block_start // 10 + 1}. Zilele {block_start + 1}–{block_start + 10}", level=2)
            block_heading.paragraph_format.page_break_before = True
        rows = []
        for item in DAILY[block_start:block_start + 10]:
            day, theme, action, output = item
            rows.append((f"Z{day}\n{business_date(int(day))}", theme, action, output))
        add_table(
            doc,
            ["Zi / dată", "Temă", "Lucru concret", "Rezultat"],
            rows,
            [Inches(0.72), Inches(1.15), Inches(3.6), Inches(1.25)],
            8.1,
        )
        if block_start == 0:
            add_callout(doc, "Gate la ziua 10", "Pagina Companii, time-study-ul și primele 10 conturi sunt gata. Dacă nu, nu începe al doilea val.", "red")
        elif block_start == 10:
            add_callout(doc, "Gate la ziua 20", "Nu executa remedieri fără acord scris, chiar dacă problema pare evidentă.", "gold")
        else:
            add_callout(doc, "Gate la ziua 30", "Rezultat minim: 1 pilot plătit sau 3 oferte active. În lipsă, pivot documentat — fără angajări și fără capex.", "red")
    doc.add_heading("11. Plan pe 13 săptămâni", level=1)
    add_table(
        doc,
        ["Săpt.", "Focus", "Prag de închidere"],
        WEEKLY,
        [Inches(0.7), Inches(2.0), Inches(4.0)],
        9.0,
    )
    doc.add_heading("Cadenta săptămânală", level=2)
    add_table(
        doc,
        ["Zi", "Bloc fondator", "Output"],
        [
            ("Luni", "90 min listă + 60 min pipeline", "10 conturi noi, next steps curate"),
            ("Marți", "2 × 60 min outbound", "10–15 atingeri"),
            ("Miercuri", "Discovery / piloți", "Date complete și oferte"),
            ("Joi", "Follow-up și închidere", "Date de decizie"),
            ("Vineri", "P&L, cash, QA, conținut", "KPI actualizați și un experiment"),
        ],
        [Inches(0.9), Inches(2.7), Inches(3.1)],
        9,
    )
    add_callout(doc, "Timp minim de vânzări", "Fondatorul blochează minimum 10 ore/săptămână pentru listă, outreach, discovery și follow-up până la 2 clienți recurenți.", "green")
    doc.add_heading("12. Plan pe 12 luni", level=1)
    add_table(
        doc,
        ["Lună", "Obiectiv", "Rezultat obligatoriu"],
        MONTHLY,
        [Inches(0.65), Inches(1.6), Inches(4.45)],
        8.8,
    )
    doc.add_heading("Porți de decizie", level=2)
    add_table(
        doc,
        ["Moment", "Continuă dacă", "Corecție dacă nu"],
        [
            ("30 zile", "1 pilot sau 3 oferte active", "Rescrie ICP și mesajul"),
            ("60 zile", "2 piloți; contribuție ≥55%", "Reprețuire / scop mai mic"),
            ("90 zile", "1–2 clienți activi", "Pivot de segment"),
            ("180 zile", "Rezultat lunar pozitiv", "Reducere cost fix / ticket mai mare"),
            ("12 luni", "Cash >45k; rezultat anual pozitiv", "Fără angajare nouă; revizie model"),
        ],
        [Inches(1.0), Inches(2.8), Inches(2.9)],
        8.8,
    )
    doc.add_heading("13. KPI, control și guvernanță", level=1)
    add_table(
        doc,
        ["KPI", "Frecvență", "Țintă", "Owner"],
        [
            ("Atingeri outbound", "Săptămânal", "30", "Fondator"),
            ("Rată răspuns", "Săptămânal", "≥10%", "Fondator"),
            ("Discovery", "Lunar", "5", "Fondator"),
            ("Piloți ofertați / executați", "Lunar", "3 / 1–2", "Fondator / atelier"),
            ("Contribuție pilot", "Per pilot", "≥55%", "Fondator"),
            ("Raport în 48h", "Lunar", "≥95%", "Atelier"),
            ("Rework", "Lunar", "<5%", "Lead tehnic"),
            ("Lei/oră tehnician", "Lunar", "≥180 lei inițial", "Fondator"),
            ("Utilizare boxe", "Lunar", "≤85%", "Fondator"),
            ("Cash", "Săptămânal", "Pozitiv; țintă buffer 45k", "Fondator"),
            ("Clienți B2B activi", "Lunar", "2 la L5; 5 la L12", "Fondator"),
        ],
        [Inches(2.55), Inches(1.25), Inches(1.65), Inches(1.25)],
        8.7,
    )
    doc.add_heading("Ședința de vineri — 30 minute", level=2)
    add_bullets(doc, [
        "5 minute: cash și facturi; 5 minute: capacitate și rework; 10 minute: funnel și oferte; 5 minute: marjă; 5 minute: experimentul următor.",
        "Orice KPI fără sursă și owner este scos din dashboard.",
        "Orice ofertă sub pragul de contribuție cere aprobare explicită și motiv documentat.",
    ])
    doc.add_heading("Dashboard-ul unic", level=2)
    p = doc.add_paragraph()
    p.add_run("Fișierul de lucru ").bold = True
    p.add_run("NOVA-PLAN-PROFITABILITATE-B2B-2026.xlsx")
    p.add_run(" conține modelul pe 12 luni, ipotezele editabile, lista de 30 prospecți, funnel-ul, planul zilnic/săptămânal/lunar și verificările.")
    doc.add_heading("14. Conformare, contractare și infrastructură", level=1)
    doc.add_heading("Checklist înainte de primul contract", level=2)
    add_table(
        doc,
        ["Domeniu", "Acțiune", "Owner"],
        [
            ("CAEN", "Confirmă CAEN Rev.3 relevant (9531) și autorizarea punctului de lucru.", "Contabil / ONRC"),
            ("TVA", "Modelează și facturează corect la cota standard curentă; planul este net de TVA.", "Contabil"),
            ("e-Factura", "Transmitere în termenul legal și reconciliere.", "Contabil"),
            ("Contract", "Comandă pilot, SLA, custodie, aprobare, răspundere, limitări.", "Avocat"),
            ("GDPR", "Informare pentru formular, CRM, poze/VIN și retenție.", "Avocat / DPO"),
            ("Mediu", "Verifică apă uzată, separator, deșeuri și autorizații locale.", "Consultant mediu"),
            ("Asigurare", "Răspundere pentru vehicule în custodie și test drive.", "Broker asigurare"),
        ],
        [Inches(1.15), Inches(4.25), Inches(1.3)],
        8.5,
    )
    doc.add_heading("Checklist tehnic înainte de trafic plătit", level=2)
    add_bullets(doc, [
        "Deploy-ul public trebuie să corespundă codului local; domeniul și DNS-ul se verifică.",
        "Formularul B2B trebuie să ajungă în inbox și în registrul intern; se testează end-to-end.",
        "PayU și avansul nu se promit până când merchant/secret/webhook sunt configurate și testate.",
        "Admin-ul, sesiunea și stocarea trebuie să funcționeze în mediul real, nu doar local.",
        "Evenimente minime: view Companii, click CTA, submit lead, rezervare începută/finalizată.",
        "Backup, retenție, export CRM și drepturile persoanelor vizate trebuie documentate.",
    ])
    add_callout(doc, "Nu este consultanță juridică sau fiscală", "Planul indică zone de verificat. Contractele, fiscalitatea, autorizațiile și obligațiile de mediu se validează cu profesioniști autorizați.", "red")
    doc.add_heading("15. Riscuri și răspunsuri", level=1)
    add_table(
        doc,
        ["Risc", "Semnal", "Răspuns"],
        [
            ("Nicio cerere B2B", "<5% răspuns după 50 atingeri", "Schimbă ICP/mesajul; nu doar volumul"),
            ("Pilot neprofitabil", "Contribuție <55%", "Reprețuire sau scop mai mic"),
            ("Supraîncărcare atelier", ">85% utilizare", "Replanifică; păstrează buffer"),
            ("Cash prea mic", "<45k buffer", "Blochează capex, costuri discreționare"),
            ("Rework", ">5%", "Stop scalare; RCA tehnic"),
            ("Ciclu corporate lung", ">45 zile fără pilot", "Revino la owner-led și divizii locale"),
            ("Promisiuni de retur", "Clientul cere garanție de cost zero", "Disclaimer; raport diagnostic"),
            ("Canal public fără decident", "3 încercări fără rol corect", "Introduceri calde / telefon centrală"),
            ("Discount eroziv", "Cerere de reducere fără volum", "Reduce scopul, nu standardul"),
            ("Dependență de fondator", ">10 clienți fără procese", "CRM, SOP și delegare după validare"),
        ],
        [Inches(2.0), Inches(2.1), Inches(2.6)],
        8.4,
    )
    doc.add_heading("Decizii care rămân blocate până la validare", level=2)
    add_bullets(doc, [
        "Al treilea tehnician, o nouă locație sau echipamente majore.",
        "Contracte B2B cu discount mare ori SLA greu de susținut.",
        "Campanii plătite ample și promisiuni de avans online.",
        "Reintroducerea VIP anual în forma economică veche.",
    ])
    doc.add_heading("16. Primele 72 de ore", level=1)
    add_kpi_cards(doc, [
        ("1", "fișă ofertă aprobată"),
        ("1", "șablon raport gata"),
        ("10", "conturi validate"),
        ("10", "mesaje pregătite"),
    ])
    doc.add_paragraph("")
    add_bullets(doc, [
        "Astăzi: aprobă oferta pilot și disclaimerul; alocă owner pentru time-study.",
        "Mâine: verifică 10 companii A și personalizează mesajele; testează formularul B2B.",
        "În 72h: trimite primul lot, sună 5 centrale și blochează două ferestre de pilot în calendar.",
        "Până vineri: actualizează workbook-ul cu Actual, Gap și următorul pas pentru fiecare cont.",
    ], numbered=True)
    add_callout(doc, "Obiectiv unic", "Primul pilot plătit este mai valoros decât încă o rundă de redesign, încă un pachet sau încă o promisiune de marketing.", "green")

    doc.add_heading("Surse principale", level=2)
    for name, description, url in SOURCES:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{name}: {description} ")
        add_hyperlink(p, "Sursă oficială", url)

    doc.add_heading("Livrabile asociate", level=2)
    add_bullets(doc, [
        "ANALIZA-OPERATIONAL-B2B-NOVA.html — audit complet și context.",
        "NOVA-PLAN-PROFITABILITATE-B2B-2026.xlsx — model financiar, prospectare și plan de execuție.",
        "companii.html — landing page B2B și formular de solicitare pilot.",
    ])
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("— Sfârșit —")
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(GOLD)

    configure_sections(doc)
    for section in doc.sections:
        add_header_footer(section)
    doc.core_properties.title = "Plan de business Nova Detailing 2026 — realist și executabil"
    doc.core_properties.subject = "Plan operațional, B2B, financiar și de vânzări"
    doc.core_properties.author = "Nova Detailing"
    doc.core_properties.keywords = "Nova Detailing, B2B, business plan, retur leasing, profitabilitate"
    doc.core_properties.comments = "Versiune refăcută după analiza operațională și B2B."
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
